#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document

def extract_text(docx_path, output_path):
    try:
        doc = Document(docx_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + '\n')
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    f.write(' | '.join(row_text) + '\n')
        print(f"Done! Content saved to {output_path}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    # File 1: Đề Ôn Thầy Trường
    extract_text('document/Đề-Ôn-Thầy-Trường-Ôn-CK.docx', 'de_on_thay_truong.txt')
    # File 2
    extract_text('document/Đề-Ôn-Thầy-Trường-Ôn-CK2.docx', 'de_on_thay_truong2.txt')
    # File 3: Trắc nghiệm nguồn mở
    extract_text('document/trắc nghiệm nguồn mở(bonus).docx', 'trac_nghiem_nguon_mo.txt')
    # File 4: Ôn tập Thầy Khuê
    extract_text('document/Ôn-tập-phát-triển-phần-mềm-mã-nguồn-mở(Thầy khuê).docx', 'on_tap_thay_khue.txt')
